import io, json, time
from datetime import datetime
from pathlib import Path

from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from logs_manager.LogsHelperManager import LogsHelperManager
from logs_manager.LogsManager import LogsManager
from zip.ZIPUtility import ZIPUtility
from zip.ZIPHelper import ZIPHelper
from data_manager.MemoryManager import MemoryManager
from tts.GTTS import GTTSService
from tts.MicrosoftEdgeTTS import MicrosoftEdgeTTS
from data_manager.DataManager import DataManager
from VoiceProcessor import VoiceProcessor
from docx import Document
from reportlab.lib.pagesizes import A4


class ZIPConvertor:
    def __init__(self, output_dir: Path):
        self.base_dir = Path(output_dir)
        self.output_dir = Path(output_dir)
        ZIPUtility.ensure_dir(self.output_dir)
        self.logger = LogsManager.get_logger("ZIPConvertor")
        self.logger.info("ZIPConvertor initialized at %s", self.base_dir)

    def split_text(self, text: str, max_chars: int = 500) -> list[str]:
        words, segs, cur = text.split(), [], ""
        for w in words:
            if len(cur) + len(w) + 1 > max_chars:
                segs.append(cur.strip()); cur = w
            else:
                cur += " " + w
        if cur:
            segs.append(cur.strip())
        return segs

    def _get_tts_service(self):
        from GUI import LANGS
        svc = (MemoryManager.get("tts_service", "google") or "").lower()
        lang = MemoryManager.get("tts_lang", "en")
        voice = MemoryManager.get("tts_voice", "female")

        if svc == "google":
            return GTTSService(lang=LANGS[lang]["gtts"]["lang"])
        elif svc == "edge":
            return MicrosoftEdgeTTS(voice=LANGS[lang]["edge"]["voices"][voice])
        else:
            raise RuntimeError(f"Unknown TTS service: {svc}")

    def export(self, text: str, fmt: str, zip_name: str = "tts_package.zip") -> Path:
        start_time = time.time()
        try:
            seg_texts = self.split_text(text, MemoryManager.get("zip_max_chars", 500))
            tts = self._get_tts_service()

            seg_files = []
            if MemoryManager.get("zip_include_segments", True):
                for i, seg in enumerate(seg_texts, start=1):
                    try:
                        seg_bytes = tts.synthesize_to_bytes(seg)
                    except Exception as e:
                        LogsHelperManager.log_error(self.logger, "SEGMENT_SYNTH_FAIL", str(e))
                        continue

                    LogsHelperManager.log_tts_request(
                        self.logger,
                        MemoryManager.get("tts_service", "unknown"),
                        MemoryManager.get("tts_lang", "en"),
                        fmt,
                        len(seg)
                    )

                    try:
                        settings = {k: MemoryManager.get(k, v) for k, v in {
                            "pitch": 0, "speed": 1.0, "volume": 1.0,
                            "echo": False, "reverb": False, "robot": False
                        }.items()}
                        processed = VoiceProcessor.process_from_memory(seg_bytes, "mp3", settings)
                        audio = DataManager.from_bytes(processed, "mp3")
                    except Exception as e:
                        LogsHelperManager.log_error(self.logger, "AUDIO_PROCESS_FAIL", str(e))
                        continue

                    try:
                        buf = io.BytesIO()
                        audio.export(buf, format=fmt)
                        seg_files.append((f"segments/segment_{i}.{fmt}", buf.getvalue()))
                    except Exception as e:
                        LogsHelperManager.log_error(self.logger, "SEGMENT_EXPORT_FAIL", str(e))
                        continue

                    LogsHelperManager.log_debug(self.logger, "SEGMENT_SYNTH",
                                                {"id": i, "bytes": len(buf.getvalue())})

            preview_len = MemoryManager.get("zip_preview_length", 200)
            preview_bytes = None
            if text.strip():
                try:
                    preview_text = text[:preview_len]
                    preview_bytes = tts.synthesize_to_bytes(preview_text)
                    LogsHelperManager.log_debug(
                        self.logger,
                        "PREVIEW_CREATED",
                        {"chars": len(preview_text), "bytes": len(preview_bytes)}
                    )
                except Exception as e:
                    LogsHelperManager.log_error(self.logger, "PREVIEW_FAIL", str(e))

            transcript_file = None
            if MemoryManager.get("zip_include_transcript", True):
                try:
                    transcript_fmt = MemoryManager.get("zip_transcript_format", "txt").lower()
                    if transcript_fmt == "txt":
                        transcript_file = ("transcript.txt", text.encode("utf-8"))
                    elif transcript_fmt == "md":
                        transcript_file = ("transcript.md", f"# Transcript\n\n{text}".encode("utf-8"))
                    elif transcript_fmt == "docx":
                        doc_buf = io.BytesIO()
                        doc = Document()
                        doc.add_heading("Transcript", 0)
                        doc.add_paragraph(text)
                        doc.save(doc_buf)
                        transcript_file = ("transcript.docx", doc_buf.getvalue())
                    elif transcript_fmt == "pdf":
                        pdf_buf = io.BytesIO()
                        doc = SimpleDocTemplate(pdf_buf, pagesize=A4)
                        styles = getSampleStyleSheet()

                        story = []
                        story.append(Paragraph("<b>Transcript</b>", styles["Title"]))
                        story.append(Spacer(1, 12))

                        for line in text.split("\n"):
                            if line.strip():
                                story.append(Paragraph(line, styles["Normal"]))
                            else:
                                story.append(Spacer(1, 12))

                        doc.build(story)
                        transcript_file = ("transcript.pdf", pdf_buf.getvalue())
                    elif transcript_fmt == "json":
                        transcript_file = (
                            "transcript.json",
                            json.dumps({"transcript": text}, indent=2).encode("utf-8")
                        )
                except Exception as e:
                    LogsHelperManager.log_error(self.logger, "TRANSCRIPT_FAIL", str(e))

            files = {}
            try:
                if seg_files and MemoryManager.get("zip_include_segments", True):
                    files["chapters.json"] = json.dumps([
                        {"id": i + 1, "file": f"segment_{i + 1}.{fmt}", "title": f"Chapter {i + 1}"}
                        for i in range(len(seg_files))
                    ], indent=2).encode("utf-8")
            except Exception as e:
                LogsHelperManager.log_error(self.logger, "CHAPTERS_JSON_FAIL", str(e))

            try:
                files["config.json"] = json.dumps({
                    "format": fmt,
                    "createdAt": datetime.now().isoformat(),
                    "service": MemoryManager.get("tts_service", "unknown"),
                    "lang": MemoryManager.get("tts_lang", "en"),
                    "voice": MemoryManager.get("tts_voice", "female"),
                    "effects": {
                        "speed": MemoryManager.get("speed", 1.0),
                        "pitch": MemoryManager.get("pitch", 0),
                        "volume": MemoryManager.get("volume", 1.0),
                        "echo": MemoryManager.get("echo", False),
                        "reverb": MemoryManager.get("reverb", False),
                        "robot": MemoryManager.get("robot", False),
                    }
                }, indent=2).encode("utf-8")
            except Exception as e:
                LogsHelperManager.log_error(self.logger, "CONFIG_JSON_FAIL", str(e))

            try:
                files["project.json"] = json.dumps({
                    "name": "Zero to Dev - Developer GUI",
                    "author": "Tuna Rasim OCAK",
                    "version": "1.0",
                    "createdAt": datetime.now().isoformat()
                }, indent=2).encode("utf-8")
            except Exception as e:
                LogsHelperManager.log_error(self.logger, "PROJECT_JSON_FAIL", str(e))

            if MemoryManager.get("zip_include_metadata", True):
                try:
                    files["metadata.json"] = json.dumps({
                        "segments": len(seg_texts),
                        "format": fmt,
                        "transcriptLength": len(text),
                        "zip_enabled": MemoryManager.get("zip_export_enabled", False),
                        "log_mode": MemoryManager.get("log_mode", "INFO")
                    }, indent=2).encode("utf-8")
                except Exception as e:
                    LogsHelperManager.log_error(self.logger, "METADATA_JSON_FAIL", str(e))

            try:
                files["events.json"] = json.dumps([{
                    "event": "convert",
                    "time": datetime.now().isoformat(),
                    "service": MemoryManager.get("tts_service", "unknown"),
                    "format": fmt,
                    "lang": MemoryManager.get("tts_lang", "en")
                }], indent=2).encode("utf-8")
            except Exception as e:
                LogsHelperManager.log_error(self.logger, "EVENTS_JSON_FAIL", str(e))

            try:
                files["readme.txt"] = (
                    "Generated by TTS Exporter\n"
                    "Includes transcript, config, chapters, metadata, preview, and segments."
                ).encode("utf-8")
            except Exception as e:
                LogsHelperManager.log_error(self.logger, "README_FAIL", str(e))

            if transcript_file:
                files[transcript_file[0]] = transcript_file[1]
            if preview_bytes:
                files[f"preview.{fmt}"] = preview_bytes

            try:
                output_path = self.output_dir / ZIPUtility.normalize_name(zip_name)
                result = ZIPHelper.create_zip(output_path, files, seg_files)
            except Exception as e:
                LogsHelperManager.log_error(self.logger, "ZIP_CREATE_FAIL", str(e))
                raise

            duration = time.time() - start_time
            LogsHelperManager.log_file_export(self.logger, str(result), result.stat().st_size)
            LogsHelperManager.log_success(self.logger, "ZIP_EXPORT", duration)
            return result

        except Exception as e:
            LogsHelperManager.log_error(self.logger, "ZIP_EXPORT", str(e))
            raise

